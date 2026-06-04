"""Ingest helpers for non-PDF source formats.

Two strategies keep the rest of the pipeline format-agnostic:

* **DjVu** is converted to PDF at upload time (djvulibre's ``ddjvu``), so every
  downstream step only ever sees a PDF.  The DjVu text layer (when present) is
  preserved, letting ``split_pdf_smart`` pull text directly and fall back to OCR
  only for image pages.
* **Plain text / .docx** are "born digital": the text already exists, so they
  skip OCR entirely — ``classify``/``extract`` read the text straight from the
  stored file.
"""

import io
import logging
import os
import subprocess
import tempfile

logger = logging.getLogger(__name__)

# Source formats whose text is already in the file (no OCR needed).
BORN_TEXT_FORMATS = {"txt", "docx"}
# All formats accepted at upload (extension without the dot).
ACCEPTED_FORMATS = {"pdf", "djvu", "txt", "docx"}


def detect_source_format(filename: str) -> str | None:
    """Map a filename to a supported source format, or None if unsupported."""
    ext = os.path.splitext(filename.lower())[1].lstrip(".")
    return ext if ext in ACCEPTED_FORMATS else None


def djvu_to_pdf(djvu_bytes: bytes, quality: int = 80, timeout: int = 1800) -> bytes:
    """Convert a (possibly multi-page) DjVu document to PDF via ddjvu.

    ``quality`` is CRITICAL.  Without it, ``ddjvu -format=pdf`` encodes every
    page LOSSLESS (Deflate TIFF), which explodes a small scanned book into a
    multi-gigabyte PDF — a 18 MB / 256-page DjVu became 2.37 GB.  Passing
    ``-quality=N`` switches the contone layers to JPEG, shrinking that ~18× to
    ~200 MB (q80) while staying perfectly legible for downstream OCR.  q80 is a
    good default for dense identification keys; lower it for pure-text scans.

    This is now invoked from the durable ``convert`` pipeline step (off the
    upload request), so the timeout can be generous — a big atlas at ~1 s/page
    still finishes in minutes.
    """
    src_path = dst_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".djvu", delete=False) as src:
            src.write(djvu_bytes)
            src_path = src.name
        dst_fd, dst_path = tempfile.mkstemp(suffix=".pdf")
        os.close(dst_fd)

        proc = subprocess.run(
            ["ddjvu", "-format=pdf", f"-quality={quality}", src_path, dst_path],
            capture_output=True, timeout=timeout,
        )
        if proc.returncode != 0:
            err = proc.stderr.decode(errors="replace")[:500]
            raise RuntimeError(f"ddjvu failed ({proc.returncode}): {err}")

        with open(dst_path, "rb") as f:
            pdf_bytes = f.read()
        if not pdf_bytes:
            raise RuntimeError("ddjvu produced an empty PDF")
        logger.info("Converted DjVu (%d bytes) -> PDF (%d bytes)", len(djvu_bytes), len(pdf_bytes))
        return pdf_bytes
    finally:
        for p in (src_path, dst_path):
            if p:
                try:
                    os.remove(p)
                except OSError:
                    pass


def extract_text_from_document(data: bytes, source_format: str) -> str:
    """Extract plain text from a born-text document (.txt / .docx)."""
    if source_format == "txt":
        return _decode_text(data)
    if source_format == "docx":
        return _extract_docx(data)
    raise ValueError(f"Unsupported born-text format: {source_format}")


def _decode_text(data: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp1251", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_docx(data: bytes) -> str:
    from docx import Document  # python-docx

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)
